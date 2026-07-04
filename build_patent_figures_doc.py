from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

base = Path(r"c:\Users\Lenovo\Desktop\ERTS\2025\newupdatedproject\Patient-monitoring-system-using-AI--main-main0.2")
figures_dir = base / "assets" / "figures"
out_doc = base / "Patent_Figures_Package_v3.docx"

# Exactly 4 pages as requested: one figure per page, no title page.
figures = [
    (
        figures_dir / "FIG1_Patient_Admission_Data_Acquisition_ZOOMED.png",
        "FIG. 1",
        "Flowchart illustrating patient admission and multimode physiological data acquisition."
    ),
    (
        figures_dir / "FIG2_ML_Feature_Engineering_Prediction_ZOOMED.png",
        "FIG. 2",
        "Flowchart illustrating feature engineering, NEWS2 scoring, and ward-specific machine-learning prediction."
    ),
    (
        figures_dir / "FIG3_Alarm_Policy_Engine_ZOOMED.png",
        "FIG. 3",
        "Flowchart illustrating context-aware alarm policy evaluation and clinical alert routing logic."
    ),
    (
        figures_dir / "FIG4_Logging_Broadcast_Discharge_ZOOMED.png",
        "FIG. 4",
        "Flowchart illustrating event logging, real-time broadcast, nurse feedback loop, and discharge processing."
    ),
]

prototype_images = [
    Path(r"C:\Users\Lenovo\Pictures\images\WhatsApp Image 2026-02-21 at 2.15.12 PM.jpeg"),
    Path(r"C:\Users\Lenovo\Pictures\images\WhatsApp Image 2026-02-21 at 2.14.27 PM.jpeg"),
    Path(r"C:\Users\Lenovo\Pictures\images\WhatsApp Image 2026-02-21 at 2.15.11 PM.jpeg"),
    Path(r"C:\Users\Lenovo\Pictures\images\WhatsApp Image 2026-02-24 at 1.58.34 PM.jpeg"),
    Path(r"C:\Users\Lenovo\Pictures\images\WhatsApp Image 2026-02-24 at 2.53.42 PM.jpeg"),
    Path(r"C:\Users\Lenovo\Pictures\images\WhatsApp Image 2026-02-21 at 2.14.27 PM (2).jpeg"),
]

doc = Document()
section = doc.sections[0]
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)

for i, (img_path, fig_label, desc) in enumerate(figures):
    if img_path.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(str(img_path), width=Inches(6.8))
    else:
        p_missing = doc.add_paragraph(f"Image not found: {img_path.name}")
        p_missing.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_missing.runs[0].font.size = Pt(10)

    # Bottom-centered patent-style label and short description.
    p_cap = doc.add_paragraph(fig_label)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = p_cap.runs[0]
    cap_run.bold = True
    cap_run.font.size = Pt(10)

    p_desc = doc.add_paragraph(desc)
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = p_desc.runs[0]
    desc_run.font.size = Pt(9)

    if i < len(figures) - 1:
        doc.add_page_break()

doc.add_page_break()

table = doc.add_table(rows=2, cols=3)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

for index, image_path in enumerate(prototype_images):
    row = index // 3
    col = index % 3
    cell = table.cell(row, col)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(2.0))

p_cap = doc.add_paragraph("FIG. 5")
p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = p_cap.runs[0]
cap_run.bold = True
cap_run.font.size = Pt(10)

p_desc = doc.add_paragraph("Prototype views of the wearable-band hardware and integrated clinical monitoring setup.")
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = p_desc.runs[0]
desc_run.font.size = Pt(9)

doc.save(out_doc)
print(f"Created: {out_doc}")
